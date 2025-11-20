import os
import json
import streamlit as st
from openai import OpenAI
from utils.parser import parse_config
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

# ---------------- API KEY ----------------
api_key = st.secrets.get("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)

# ---------------- PAGE CONFIG ----------------
st.set_page_config(page_title="NetDoc AI", layout="wide")

# ---------------- LOGO ----------------
st.image("logo.png", width=180)

st.title("⚡ NetDoc AI — Autonomous Network Documentation Engine")
st.write("Upload configs → AI generates documentation, auditing, topology & exports.")

# ---------------- FILE UPLOAD ----------------
uploaded_files = st.file_uploader(
    "Upload one or more config files",
    type=["txt", "cfg", "log"],
    accept_multiple_files=True
)

# ---------------- HELPER FUNCTIONS ----------------

def generate_pdf(text):
    pdf_bytes = "/tmp/netdoc.pdf"
    c = canvas.Canvas(pdf_bytes, pagesize=letter)
    width, height = letter

    y = height - 50
    c.setFont("Helvetica", 10)

    for line in text.split("\n"):
        if y < 50:
            c.showPage()
            y = height - 50
            c.setFont("Helvetica", 10)

        c.drawString(40, y, line[:120])  # Protects overflow
        y -= 15

    c.save()

    with open(pdf_bytes, "rb") as f:
        return f.read()


def generate_docx(text):
    doc = Document()
    doc.add_heading("NetDoc AI — Report", level=1)

    for line in text.split("\n"):
        doc.add_paragraph(line)

    path = "/tmp/netdoc.docx"
    doc.save(path)
    with open(path, "rb") as f:
        return f.read()


def generate_html(text):
    html = f"""
    <html><body>
    <h2>NetDoc AI — Network Documentation Report</h2>
    <pre>{text}</pre>
    </body></html>
    """
    return html.encode("utf-8")


def ai_topology(data):
    prompt = f"""
    You are a network topology generator.

    Based ONLY on this parsed config data, produce an ASCII topology diagram.

    Data:
    {json.dumps(data, indent=2)}
    """

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )

    return response.choices[0].message["content"]


def ai_security_audit(data):
    prompt = f"""
    You are a senior network security engineer.

    Perform a configuration security audit.

    Output sections:
    - 🔐 Critical Issues
    - ⚠️ Warnings
    - 📝 Best Practices
    - 🛠 Remediation Steps

    Data:
    {json.dumps(data, indent=2)}
    """

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )

    return response.choices[0].message["content"]


# ---------------- PROCESS CONFIGS ----------------
if uploaded_files and st.button("Generate Report"):
    combined = ""
    for f in uploaded_files:
        combined += f"\n\n# FILE: {f.name}\n"
        combined += f.read().decode("utf-8")

    with st.spinner("Analyzing your configs..."):
        parsed = parse_config(combined)

    st.success("Report generated!")

    st.session_state["parsed"] = parsed
    st.session_state["json_md"] = json.dumps(parsed, indent=2)

# ---------------- TABS ----------------
tab1, tab2, tab3, tab4 = st.tabs(["📄 Overview", "🛡 Security Audit", "📡 Topology", "📤 Export"])

# ---------------- OVERVIEW TAB ----------------
with tab1:
    if "parsed" in st.session_state:
        st.subheader("Parsed Configuration Details")
        st.json(st.session_state["parsed"])
    else:
        st.info("Upload configs to generate the report.")

# ---------------- SECURITY AUDIT TAB ----------------
with tab2:
    if "parsed" in st.session_state:
        if st.button("Run Security Audit"):
            with st.spinner("Running audit..."):
                audit_text = ai_security_audit(st.session_state["parsed"])
            st.session_state["audit"] = audit_text
        if "audit" in st.session_state:
            st.subheader("Security Audit Report")
            st.markdown(st.session_state["audit"])
    else:
        st.info("Generate a report first.")

# ---------------- TOPOLOGY TAB ----------------
with tab3:
    if "parsed" in st.session_state:
        if st.button("Generate Topology Diagram"):
            with st.spinner("Building topology diagram..."):
                topo = ai_topology(st.session_state["parsed"])
            st.session_state["topo"] = topo

        if "topo" in st.session_state:
            st.subheader("ASCII Topology Map")
            st.code(st.session_state["topo"])
    else:
        st.info("Generate a report first.")

# ---------------- EXPORT TAB ----------------
with tab4:
    if "json_md" in st.session_state:
        st.subheader("Export Formats")

        md = st.session_state["json_md"]

        # PDF
        pdf = generate_pdf(md)
        st.download_button(
            "📄 Download PDF",
            data=pdf,
            file_name="NetDoc_Report.pdf",
            mime="application/pdf"
        )

        # DOCX
        docx = generate_docx(md)
        st.download_button(
            "📝 Download DOCX",
            data=docx,
            file_name="NetDoc_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # HTML
        html = generate_html(md)
        st.download_button(
            "🌐 Download HTML",
            data=html,
            file_name="NetDoc_Report.html",
            mime="text/html"
        )
    else:
        st.info("Generate a report to export.")
