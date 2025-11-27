# ============================================================
# main.py — Security Audit + Topology Engine + Export Engine
# Complete Part 2 + Part 3 + Part 4 in one file
# ============================================================

import json
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document

# ============================
# SECURITY AUDIT ENGINE (C-27)
# ============================

def run_security_audit(parsed):
    issues = []

    # --- Hostname check ---
    if not parsed["device_summary"].get("hostname"):
        issues.append("❌ Hostname missing — configure a unique hostname.")

    # --- Password issues ---
    weak_pw_patterns = ["password", "cisco", "admin", "1234", "12345"]
    config_text = json.dumps(parsed)

    for p in weak_pw_patterns:
        if p.lower() in config_text.lower():
            issues.append(f"❌ Weak password detected: '{p}'")

    # --- VLAN issues ---
    for v in parsed["vlans"]:
        if not v.get("ports"):
            issues.append(f"⚠ VLAN {v['vlan_id']} has no active ports.")

    # --- STP issues ---
    if "spanning-tree portfast" not in config_text:
        issues.append("⚠ STP PortFast appears missing on access ports.")

    # --- Default Route ---
    if not parsed["routing_summary"].get("default_route"):
        issues.append("⚠ No default route found.")

    return issues if issues else ["✅ No major issues detected. Good job!"]


# ============================
# TOPOLOGY ENGINE (C-28)
# Mermaid diagram output
# ============================

def generate_topology_mermaid(parsed):
    neighbors = parsed.get("neighbors", [])
    if not neighbors:
        return "No topology data found."

    mermaid = ["graph LR;"]
    for n in neighbors:
        local = n["local_interface"].replace("/", "_")
        nbr = n["neighbor_interface"].replace("/", "_")
        dev = n["neighbor_device"]
        mermaid.append(f'    {local}["{local}"] -- {dev} --> {nbr}["{nbr}"];')

    return "\n".join(mermaid)


# ============================
# EXPORT ENGINE (C-29)
# PDF + DOCX + HTML
# ============================

# ---- HTML EXPORT ----
def export_html(parsed):
    return f"""
    <html>
    <head><title>NetDoc AI Export</title></head>
    <body style="font-family: Helvetica; margin: 40px;">
    <h1>NetDoc AI — Documentation Export</h1>
    <pre>{json.dumps(parsed, indent=2)}</pre>
    </body>
    </html>
    """


# ---- PDF EXPORT ----
def export_pdf(parsed):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    c.setFont("Helvetica", 10)

    text = json.dumps(parsed, indent=2)
    y = 750

    for line in text.split("\n"):
        if y < 40:  # new page
            c.showPage()
            c.setFont("Helvetica", 10)
            y = 750

        c.drawString(40, y, line)
        y -= 14

    c.save()
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes


# ---- DOCX EXPORT ----
def export_docx(parsed):
    doc = Document()
    doc.add_heading('NetDoc AI — Documentation Export', level=1)
    doc.add_paragraph(json.dumps(parsed, indent=2))

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ---- EXPORT ALL ----
def export_all_formats(parsed):
    return {
        "html": export_html(parsed),
        "pdf": export_pdf(parsed),
        "docx": export_docx(parsed),
    }
